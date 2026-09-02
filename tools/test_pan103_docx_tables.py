"""PAN-103 -- DOCX real table-grid extraction via docx2python.

parse_docx previously read only w:t text nodes -- table-cell text was
included but completely flattened, with no row/column boundary at all.
Verified against a real 41-table Keystone document (PANTA Canonical
Investment Case v1.1): a 6-row, 2-column "what changed" table came out as
one undifferentiated paragraph before this fix. Tests here use a small
synthetic document built the same way python-docx itself would build a
real one, not a hand-rolled minimal package (docx2python needs a genuinely
valid OOXML package, same lesson already learned for PPTX in PAN-101).
"""

import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from tools.extract_v2 import parse_docx  # noqa: E402


class PAN103DocxTableTests(unittest.TestCase):
    def setUp(self):
        self.temporary = tempfile.TemporaryDirectory()
        self.path = Path(self.temporary.name) / "report.docx"

        document = Document()
        document.add_paragraph("Executive summary")
        document.add_paragraph("Revenue was EUR 20m in FY2025A.")

        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text, table.cell(0, 1).text = "Metric", "Value"
        table.cell(1, 0).text, table.cell(1, 1).text = "Revenue", "EUR 20m"
        table.cell(2, 0).text, table.cell(2, 1).text = "EBITDA margin", "13.8%"

        document.add_paragraph("Management expects growth in FY2026E.")
        document.save(str(self.path))

    def tearDown(self):
        self.temporary.cleanup()

    def _body(self) -> str:
        chunks = parse_docx(self.path)
        return "\n\n".join(c.body for c in chunks)

    def test_table_extracted_as_structured_grid_not_flattened_text(self):
        body = self._body()
        self.assertIn("| Metric | Value |", body)
        self.assertIn("| Revenue | EUR 20m |", body)
        self.assertIn("| EBITDA margin | 13.8% |", body)

    def test_paragraphs_before_and_after_table_still_captured(self):
        body = self._body()
        self.assertIn("Executive summary", body)
        self.assertIn("Revenue was EUR 20m in FY2025A.", body)
        self.assertIn("Management expects growth in FY2026E.", body)

    def test_falls_back_to_paragraph_text_when_docx2python_unavailable(self):
        with patch.dict(sys.modules, {"docx2python": None}):
            chunks = parse_docx(self.path)
        body = "\n\n".join(c.body for c in chunks)
        # No table grid available in the fallback, but no data loss either --
        # cell text still comes through as plain paragraph-shaped text.
        self.assertIn("Revenue", body)
        self.assertIn("EUR 20m", body)
        self.assertNotIn("| Metric | Value |", body)


if __name__ == "__main__":
    unittest.main()
