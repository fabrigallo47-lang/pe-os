"""Native synthetic documents shared by connected API tests and the isolated UI lab."""
import hashlib
import io
import json
import wave
from pathlib import Path


def pdf_bytes():
    # Small valid two-page PDF; no optional fixture-generation dependency.
    texts = ["Synthetic company document", "Reported performance requires independent verification."]
    objects = [b"<< /Type /Catalog /Pages 2 0 R >>", b"<< /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >>"]
    objects += [f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents {6+i} 0 R >>".encode() for i in range(2)]
    objects += [b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"]
    for text in texts:
        stream = f"BT /F1 16 Tf 48 720 Td ({text}) Tj ET".encode()
        objects.append(f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream")
    result, offsets = b"%PDF-1.4\n", [0]
    for i, obj in enumerate(objects, 1):
        offsets.append(len(result))
        result += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref = len(result)
    result += f"xref\n0 {len(offsets)}\n0000000000 65535 f \n".encode()
    result += b"".join(f"{offset:010} 00000 n \n".encode() for offset in offsets[1:])
    result += f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    return result


def fixture_files():
    from openpyxl import Workbook
    from docx import Document
    book = Workbook()
    sheet = book.active
    sheet.title = "Financing Plan"
    sheet.append(["Metric", "Value EURm", "Formula"])
    sheet.append(["Raise size", 5, "=B2*2"])
    sheet.append(["Runway months", 18, None])
    spreadsheet = io.BytesIO()
    book.save(spreadsheet)
    doc = Document()
    doc.add_paragraph("Synthetic opening note")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "Topic", "Evidence"
    table.cell(1, 0).text, table.cell(1, 1).text = "Deployment", "Confirmed"
    doc.add_paragraph("Closing context")
    word = io.BytesIO()
    doc.save(word)
    recording = io.BytesIO()
    with wave.open(recording, "wb") as audio:
        audio.setparams((1, 2, 8000, 0, "NONE", "not compressed"))
        audio.writeframes(b"\x00\x00" * 8000 * 4)
    return {
        "company.pdf": pdf_bytes(),
        "financing.xlsx": spreadsheet.getvalue(),
        "call.srt": b"1\n00:00:01,000 --> 00:00:02,000\nSynthetic introduction\n\n2\n00:03:12,000 --> 00:03:29,000\nWe know the workflow well. Senior delivery hires are still planned.\n",
        "customer.txt": b"Synthetic customer reference\nDeployment is confirmed.\nIndependent performance evidence remains outstanding.\n<script>alert('source text')</script>\n",
        "note.docx": word.getvalue(),
        "recording.wav": recording.getvalue(),
    }


def build_fixture_case(root: Path, case_id="CASE-1"):
    inbox = root / "vault" / "inbox"
    inbox.mkdir(parents=True)
    specs = [
        ("SRC-2", "CL-2", "company.pdf", "p2", "Reported performance requires independent verification."),
        ("SRC-6", "CL-8", "financing.xlsx", "financing.xlsx::'Financing Plan'!B2:C2", ""),
        ("SRC-3", "CL-3", "call.srt", "call.srt::cue:2:00:03:12.000-->00:03:29.000", "We know the workflow well. Senior delivery hires are still planned."),
        ("SRC-1", "CL-1", "customer.txt", "customer.txt::lines:2-2", "Deployment is confirmed."),
        ("SRC-WORD", "CL-WORD", "note.docx", "note.docx::blocks:2-2", ""),
        ("SRC-AUDIO", "CL-AUDIO", "recording.wav", "00:00:01.000–00:00:02.000", ""),
    ]
    files, records, citations = fixture_files(), [], []
    for source, claim, filename, locator, quote in specs:
        data = files[filename]
        (inbox / filename).write_bytes(data)
        version = "sha256:" + hashlib.sha256(data).hexdigest()
        envelope = {"schema_version": "panta.source-envelope/1.0", "case_id": case_id, "source_id": source,
                    "source_version_id": version, "original_filename": filename, "stored_filename": filename}
        records.append({"case_id": case_id, "source_envelope": envelope})
        citations.append({"sourceId": source, "claimId": claim, "sourceVersionId": version, "locator": locator,
                          "filename": filename, "verbatimOrLosslessSpan": quote})
    (inbox / ".ingest-manifest.json").write_text(json.dumps({"items": records}))
    bundle = root / "cases" / case_id
    bundle.mkdir(parents=True)
    claims = [{"claim_id": item["claimId"], "source_id": item["sourceId"], "source_version_id": item["sourceVersionId"],
               "locator": item["locator"], "verbatim_or_lossless_span": item["verbatimOrLosslessSpan"]} for item in citations]
    (bundle / "claims.json").write_text(json.dumps(claims))
    return {"caseId": case_id, "citations": citations, "records": records, "claims": claims}
